from docx import Document


class DocumentService:
    def __init__(self, logger):
        self.logger = logger
        self.document = None
    
    def create_document(self):
        self.logger.info("Criando documento...")
        self.document = Document()

    def insert_select_results(self, results):
        self.logger.info("Inserindo selects no documento...")
        self.document.add_heading("Selects", level=1)
        for i, result in enumerate(results, start=1):
            self.logger.info(f"Inserindo select {i} no documento...")
            self.document.add_heading(f"Objetivo: {result[0]}", level=2)
            self.document.add_paragraph(f"Select realizado: {result[1]}")
            self.document.add_paragraph(f"Tempo de execução: {format((result[2] * 1000), '.3f')} milissegundos")
            self.document.add_paragraph(f"Quantidade de dados retornados: {(result[3])}")
            self.document.add_paragraph(f"Dados retornados:")
            if i == len(results):
                self.logger.info("O de cpf da mais de 900000 registros, não será inserido no documento")
                self.logger.info("Sucesso!")
                break
            table = self.document.add_table(1, len(result[5][0]))
            #get last in results
            table.style = 'Table Grid'
            heading_cells = table.rows[0].cells
            for i, label in enumerate(result[4]):
                heading_cells[i].text = label
            for i, row in enumerate(result[5]):
                cells = table.add_row().cells
                for j in range(len(cells)):
                    cells[j].text = str(row[j])

            self.document.add_page_break()

            
        self.logger.info("Selects inseridos no documento!")
    def save(self, name):
        self.logger.info(f"Salvando documento {name}...")
        self.document.save(name)
        self.logger.info("Documento salvo com sucesso!")